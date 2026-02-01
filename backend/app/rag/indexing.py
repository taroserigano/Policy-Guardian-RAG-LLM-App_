"""
Document indexing pipeline: extract text, chunk, embed, and store in Pinecone.
Handles PDF, TXT, and Word (.docx) files.
"""
from typing import List, Dict, Any, BinaryIO
import io
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pinecone import Pinecone, ServerlessSpec
import pypdf

from app.core.config import get_settings
from app.core.logging import get_logger
from app.rag.embeddings import get_default_embeddings

# Optional: python-docx for Word document support
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

settings = get_settings()
logger = get_logger(__name__)

# Initialize Pinecone client
pc = Pinecone(api_key=settings.pinecone_api_key)

# Pinecone metadata limits
MAX_METADATA_SIZE = 40000  # 40KB limit for metadata text


def sanitize_text_for_pinecone(text: str) -> str:
    """
    Sanitize text for Pinecone storage by:
    - Removing null bytes and other problematic characters
    - Ensuring valid UTF-8 encoding
    - Truncating to fit metadata size limits
    
    Args:
        text: Raw text to sanitize
    
    Returns:
        Sanitized text safe for Pinecone metadata
    """
    if not text:
        return ""
    
    # Remove null bytes and other control characters (except newlines and tabs)
    sanitized = ""
    for char in text:
        if char in ('\n', '\r', '\t') or (ord(char) >= 32 and ord(char) != 127):
            sanitized += char
    
    # Ensure valid UTF-8 (replace invalid sequences)
    sanitized = sanitized.encode('utf-8', errors='replace').decode('utf-8')
    
    # Truncate if too long for metadata (leave some room for other fields)
    if len(sanitized.encode('utf-8')) > MAX_METADATA_SIZE:
        # Truncate to fit
        while len(sanitized.encode('utf-8')) > MAX_METADATA_SIZE:
            sanitized = sanitized[:-100]
        sanitized += "... [truncated]"
    
    return sanitized.strip()


def ensure_index_exists() -> None:
    """
    Create Pinecone index if it doesn't exist.
    Uses serverless spec for cost efficiency.
    """
    index_name = settings.pinecone_index_name
    
    if index_name not in pc.list_indexes().names():
        logger.info(f"Creating Pinecone index: {index_name}")
        pc.create_index(
            name=index_name,
            dimension=settings.embed_dim,
            metric="cosine",
            spec=ServerlessSpec(
                cloud=settings.pinecone_cloud,
                region=settings.pinecone_region
            )
        )
        logger.info(f"Index {index_name} created successfully")
    else:
        logger.info(f"Index {index_name} already exists")


def get_pinecone_index():
    """Get Pinecone index instance."""
    ensure_index_exists()
    return pc.Index(settings.pinecone_index_name)


def extract_text_from_pdf(file_path: str) -> tuple[str, List[Dict[str, Any]]]:
    """
    Extract text from PDF file with page numbers.
    
    Args:
        file_path: Path to PDF file
    
    Returns:
        Tuple of (full_text, pages_metadata)
        pages_metadata contains page numbers and text per page
    """
    try:
        # Use pypdf directly instead of langchain loader
        full_text = ""
        pages_metadata = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            for i, page in enumerate(pdf_reader.pages):
                page_num = i + 1
                page_text = page.extract_text() or ""
                # Sanitize extracted text
                page_text = sanitize_text_for_pinecone(page_text)
                full_text += page_text + "\n\n"
                pages_metadata.append({
                    "page_number": page_num,
                    "text": page_text
                })
        
        return full_text, pages_metadata
    
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from TXT file.
    
    Args:
        file_path: Path to TXT file
    
    Returns:
        Text content
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        raise


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from Word document (.docx).
    
    Args:
        file_path: Path to Word document
    
    Returns:
        Text content extracted from all paragraphs and tables
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for Word document support. "
            "Install it with: pip install python-docx"
        )
    
    try:
        doc = DocxDocument(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        full_text = "\n\n".join(paragraphs)
        logger.info(f"Extracted {len(paragraphs)} paragraphs/rows from Word document")
        return full_text
    
    except Exception as e:
        logger.error(f"Error extracting text from Word document: {e}")
        raise


def chunk_text(text: str) -> List[str]:
    """
    Split text into chunks using RecursiveCharacterTextSplitter.
    Uses section-aware separators for better document structure preservation.
    
    Args:
        text: Full document text
    
    Returns:
        List of text chunks
    """
    # Section-aware separators - prioritize keeping sections together
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=[
            "\n\n\n",  # Multiple newlines (section breaks)
            "\n\n",    # Paragraph breaks
            "\n",      # Line breaks
            ". ",      # Sentence boundaries
            "; ",      # Clause boundaries
            ", ",      # Phrase boundaries
            " ",       # Word boundaries
            ""         # Character fallback
        ],
        keep_separator=True  # Keep separators for context
    )
    chunks = splitter.split_text(text)
    logger.info(f"Split text into {len(chunks)} chunks")
    return chunks


def index_document(
    doc_id: str,
    filename: str,
    file_path: str,
    content_type: str
) -> Dict[str, Any]:
    """
    Complete indexing pipeline: extract, chunk, embed, and upsert to Pinecone.
    
    Args:
        doc_id: Unique document ID (UUID)
        filename: Original filename
        file_path: Path to file on disk
        content_type: MIME type (application/pdf, text/plain, or application/vnd.openxmlformats-officedocument.wordprocessingml.document)
    
    Returns:
        Dictionary with preview_text and chunk_count
    """
    logger.info(f"Starting indexing for document: {filename}")
    
    # Step 1: Extract text based on content type
    page_mapping = None
    
    if content_type == "application/pdf":
        full_text, pages_metadata = extract_text_from_pdf(file_path)
        page_mapping = {i: meta["page_number"] for i, meta in enumerate(pages_metadata)}
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        full_text = extract_text_from_docx(file_path)
    else:
        # Default to plain text
        full_text = extract_text_from_txt(file_path)
    
    # Sanitize full text
    full_text = sanitize_text_for_pinecone(full_text)
    
    if not full_text.strip():
        logger.warning(f"No text extracted from document: {filename}")
        raise ValueError(f"No readable text content found in {filename}")
    
    # Step 2: Chunk text
    chunks = chunk_text(full_text)
    
    if not chunks:
        logger.warning(f"No chunks created from document: {filename}")
        raise ValueError(f"Could not create text chunks from {filename}")
    
    # Step 3: Generate embeddings
    logger.info(f"Generating embeddings for {len(chunks)} chunks...")
    embeddings = get_default_embeddings().embed_documents(chunks)
    
    # Step 4: Prepare vectors for Pinecone
    vectors = []
    for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        vector_id = f"{doc_id}:{i}"
        
        # Sanitize chunk text for metadata storage
        sanitized_chunk = sanitize_text_for_pinecone(chunk)
        
        # Build metadata (ensure all values are safe for Pinecone)
        metadata = {
            "doc_id": str(doc_id),
            "filename": sanitize_text_for_pinecone(filename)[:500],  # Limit filename length
            "source_type": str(content_type),
            "chunk_index": int(i),
            "text": sanitized_chunk  # Store sanitized text in metadata for retrieval
        }
        
        # Add page number for PDFs
        if page_mapping:
            # Estimate which page this chunk came from
            # Simple heuristic: divide by number of chunks
            estimated_page = min(
                max(1, int((i / len(chunks)) * len(page_mapping)) + 1),
                len(page_mapping)
            )
            metadata["page_number"] = int(estimated_page)
        
        vectors.append({
            "id": vector_id,
            "values": embedding,
            "metadata": metadata
        })
    
    # Step 5: Upsert to Pinecone
    logger.info(f"Upserting {len(vectors)} vectors to Pinecone...")
    index = get_pinecone_index()
    index.upsert(vectors=vectors, namespace="")
    
    logger.info(f"Successfully indexed document {filename}")
    
    # Return preview text (first 500 chars)
    preview_text = full_text[:500] if len(full_text) > 500 else full_text
    
    return {
        "preview_text": preview_text,
        "chunk_count": len(chunks)
    }


def delete_document_from_index(doc_id: str) -> None:
    """
    Delete all vectors for a document from Pinecone.
    
    Args:
        doc_id: Document UUID
    """
    try:
        index = get_pinecone_index()
        
        # Query to get all vector IDs for this document
        # Pinecone doesn't support prefix deletion directly, so we fetch and delete
        logger.info(f"Deleting vectors for document: {doc_id}")
        
        # Use delete with filter
        index.delete(filter={"doc_id": {"$eq": doc_id}})
        
        logger.info(f"Deleted all vectors for document {doc_id}")
    
    except Exception as e:
        logger.error(f"Error deleting document from index: {e}")
        raise

