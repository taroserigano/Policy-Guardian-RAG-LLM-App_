"""
Generate PDF and Word versions of sample policies and upload them to the RAG system.
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
import requests

# Add backend to path
sys.path.insert(0, str(Path(__file__).parent))

BACKEND_URL = "http://localhost:8000"
SAMPLE_DOCS_DIR = Path(__file__).parent / "sample_docs"


def create_remote_work_policy_pdf():
    """Create a PDF version of the remote work policy."""
    output_path = SAMPLE_DOCS_DIR / "remote_work_policy.pdf"
    
    # Read the text version
    txt_path = SAMPLE_DOCS_DIR / "remote_work_policy.txt"
    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create PDF
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18,
    )
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#2563eb'),
        spaceAfter=30,
        alignment=1  # Center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=12,
        spaceBefore=12,
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=10,
        spaceAfter=6,
    )
    
    # Parse and add content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            elements.append(Spacer(1, 0.1*inch))
            continue
        
        # Title (first line or all caps with "POLICY" in it)
        if 'REMOTE WORK POLICY' in line and len(line) < 50:
            elements.append(Paragraph(line, title_style))
            elements.append(Spacer(1, 0.2*inch))
        # Section headings (numbered or all caps short lines)
        elif (line[0].isdigit() and '.' in line[:5]) or (line.isupper() and len(line) < 80):
            elements.append(Paragraph(line, heading_style))
        # Regular text
        else:
            elements.append(Paragraph(line, body_style))
    
    # Build PDF
    doc.build(elements)
    print(f"✓ Created PDF: {output_path}")
    return output_path


def create_employee_leave_policy_docx():
    """Create a Word document version of the employee leave policy."""
    output_path = SAMPLE_DOCS_DIR / "employee_leave_policy.docx"
    
    # Read the text version
    txt_path = SAMPLE_DOCS_DIR / "employee_leave_policy.txt"
    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Parse and add content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Title
        if 'EMPLOYEE LEAVE POLICY' in line and len(line) < 50:
            heading = doc.add_heading(line, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Major section headings (numbered)
        elif line[0].isdigit() and '.' in line[:5] and line[1].isdigit():
            doc.add_heading(line, level=1)
        # Subsection headings (numbered with more decimals)
        elif line[0].isdigit() and '.' in line[:5]:
            doc.add_heading(line, level=2)
        # Metadata lines
        elif any(x in line for x in ['Organization:', 'Version:', 'Effective Date:', 'Classification:', 'Important Note']):
            p = doc.add_paragraph(line)
            if 'Important Note' in line:
                p.runs[0].bold = True
        # Regular text
        else:
            doc.add_paragraph(line)
    
    # Save document
    doc.save(output_path)
    print(f"✓ Created Word doc: {output_path}")
    return output_path


def create_data_privacy_policy_pdf():
    """Create a PDF version of the data privacy policy."""
    output_path = SAMPLE_DOCS_DIR / "data_privacy_policy.pdf"
    
    # Read the text version
    txt_path = SAMPLE_DOCS_DIR / "data_privacy_policy.txt"
    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create PDF
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18,
    )
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#7c3aed'),
        spaceAfter=20,
        alignment=1
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.HexColor('#5b21b6'),
        spaceAfter=10,
        spaceBefore=10,
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=10,
        spaceAfter=6,
    )
    
    # Parse content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            elements.append(Spacer(1, 0.1*inch))
            continue
        
        if 'DATA PRIVACY' in line and 'POLICY' in line and len(line) < 80:
            elements.append(Paragraph(line, title_style))
            elements.append(Spacer(1, 0.2*inch))
        elif (line[0].isdigit() and '.' in line[:5]) or (len(line) < 80 and line.isupper()):
            elements.append(Paragraph(line, heading_style))
        else:
            elements.append(Paragraph(line, body_style))
    
    doc.build(elements)
    print(f"✓ Created PDF: {output_path}")
    return output_path


def upload_file_to_rag(file_path):
    """Upload a file to the RAG system."""
    try:
        with open(file_path, 'rb') as f:
            files = {'file': (file_path.name, f)}
            response = requests.post(
                f"{BACKEND_URL}/api/docs/upload",
                files=files,
                timeout=120
            )
        
        if response.status_code in [200, 201]:
            data = response.json()
            print(f"  ✅ Uploaded: {file_path.name} (ID: {data.get('id', 'unknown')})")
            return True
        else:
            print(f"  ❌ Failed to upload {file_path.name}: {response.status_code}")
            print(f"     Response: {response.text}")
            return False
    except Exception as e:
        print(f"  ❌ Error uploading {file_path.name}: {e}")
        return False


def main():
    print("="*60)
    print("  GENERATING SAMPLE PDF AND WORD DOCUMENTS")
    print("="*60)
    print()
    
    # Create documents
    print("Creating documents...")
    files_to_upload = []
    
    try:
        pdf1 = create_remote_work_policy_pdf()
        files_to_upload.append(pdf1)
    except Exception as e:
        print(f"❌ Error creating Remote Work PDF: {e}")
    
    try:
        docx1 = create_employee_leave_policy_docx()
        files_to_upload.append(docx1)
    except Exception as e:
        print(f"❌ Error creating Employee Leave Word doc: {e}")
    
    try:
        pdf2 = create_data_privacy_policy_pdf()
        files_to_upload.append(pdf2)
    except Exception as e:
        print(f"❌ Error creating Data Privacy PDF: {e}")
    
    print()
    print("="*60)
    print(f"  UPLOADING {len(files_to_upload)} DOCUMENTS TO RAG SYSTEM")
    print("="*60)
    print()
    
    # Check if backend is running
    try:
        response = requests.get(f"{BACKEND_URL}/health", timeout=5)
        if response.status_code != 200:
            print("⚠️  Backend server not responding correctly")
            print(f"   Make sure the server is running on {BACKEND_URL}")
            print(f"   You can still find the generated files in: {SAMPLE_DOCS_DIR}")
            return
    except requests.exceptions.RequestException:
        print(f"❌ Cannot connect to backend at {BACKEND_URL}")
        print("   Make sure the server is running with: python -m uvicorn app.main:app --host 0.0.0.0 --port 8000")
        print()
        print(f"✓ Generated files are saved in: {SAMPLE_DOCS_DIR}")
        print("  You can upload them manually through the UI once the server is running.")
        return
    
    # Upload files
    success_count = 0
    for file_path in files_to_upload:
        if upload_file_to_rag(file_path):
            success_count += 1
    
    print()
    print("="*60)
    print(f"  COMPLETE: {success_count}/{len(files_to_upload)} files uploaded successfully")
    print("="*60)
    print()
    print("📁 Generated files location:")
    for file_path in files_to_upload:
        print(f"   - {file_path.name}")
    print()
    print("🌐 Test these files in the UI:")
    print("   - Upload page: http://localhost:5174/upload")
    print("   - Chat page: http://localhost:5174/chat")
    print()


if __name__ == "__main__":
    main()
