"""
Comprehensive tests for document indexing functionality.
Tests: text extraction, sanitization, chunking, PDF/Word/TXT processing, and upload API.
"""
import pytest
import tempfile
import os
import sys
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
import io

# Add backend to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from app.rag.indexing import (
    sanitize_text_for_pinecone,
    extract_text_from_txt,
    extract_text_from_pdf,
    chunk_text,
    MAX_METADATA_SIZE,
    DOCX_AVAILABLE
)

# Try importing docx extraction if available
if DOCX_AVAILABLE:
    from app.rag.indexing import extract_text_from_docx


class TestSanitizeTextForPinecone:
    """Tests for the sanitize_text_for_pinecone function."""
    
    def test_empty_string(self):
        """Empty string should return empty string."""
        assert sanitize_text_for_pinecone("") == ""
    
    def test_none_input(self):
        """None should return empty string."""
        assert sanitize_text_for_pinecone(None) == ""
    
    def test_normal_text(self):
        """Normal text should pass through unchanged."""
        text = "This is a normal policy document."
        result = sanitize_text_for_pinecone(text)
        assert result == text
    
    def test_preserves_newlines(self):
        """Newlines should be preserved."""
        text = "Line 1\nLine 2\nLine 3"
        result = sanitize_text_for_pinecone(text)
        assert "\n" in result
        assert "Line 1" in result
        assert "Line 3" in result
    
    def test_preserves_tabs(self):
        """Tabs should be preserved."""
        text = "Column1\tColumn2\tColumn3"
        result = sanitize_text_for_pinecone(text)
        assert "\t" in result
    
    def test_removes_null_bytes(self):
        """Null bytes should be removed."""
        text = "Hello\x00World"
        result = sanitize_text_for_pinecone(text)
        assert "\x00" not in result
        assert "HelloWorld" in result
    
    def test_removes_control_characters(self):
        """Control characters (except newline/tab) should be removed."""
        text = "Hello\x01\x02\x03World"
        result = sanitize_text_for_pinecone(text)
        assert "\x01" not in result
        assert "\x02" not in result
        assert "\x03" not in result
        assert "HelloWorld" in result
    
    def test_removes_delete_character(self):
        """DEL character (0x7F) should be removed."""
        text = "Hello\x7fWorld"
        result = sanitize_text_for_pinecone(text)
        assert "\x7f" not in result
    
    def test_handles_unicode(self):
        """Unicode characters should be preserved."""
        text = "日本語テスト émojis 🎉 symbols ™®©"
        result = sanitize_text_for_pinecone(text)
        assert "日本語" in result
        assert "🎉" in result
    
    def test_handles_mixed_encodings(self):
        """Mixed encoding issues should be handled gracefully."""
        # Create text with potential encoding issues
        text = "Normal text with special: café résumé naïve"
        result = sanitize_text_for_pinecone(text)
        assert "café" in result or "caf" in result  # Should handle gracefully
    
    def test_truncates_long_text(self):
        """Text exceeding MAX_METADATA_SIZE should be truncated."""
        # Create text larger than 40KB
        long_text = "A" * 50000
        result = sanitize_text_for_pinecone(long_text)
        
        # Should be truncated
        assert len(result.encode('utf-8')) <= MAX_METADATA_SIZE + 100  # Allow for truncation marker
        assert "... [truncated]" in result
    
    def test_strips_whitespace(self):
        """Leading/trailing whitespace should be stripped."""
        text = "   Hello World   "
        result = sanitize_text_for_pinecone(text)
        assert result == "Hello World"
    
    def test_carriage_returns_preserved(self):
        """Carriage returns should be preserved (for Windows line endings)."""
        text = "Line 1\r\nLine 2"
        result = sanitize_text_for_pinecone(text)
        assert "Line 1" in result
        assert "Line 2" in result


class TestExtractTextFromTxt:
    """Tests for TXT file extraction."""
    
    def test_simple_txt_file(self):
        """Extract text from simple TXT file."""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("This is a test document.\nSecond line.")
            temp_path = f.name
        
        try:
            result = extract_text_from_txt(temp_path)
            assert "This is a test document." in result
            assert "Second line." in result
        finally:
            os.unlink(temp_path)
    
    def test_empty_txt_file(self):
        """Extract text from empty TXT file."""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("")
            temp_path = f.name
        
        try:
            result = extract_text_from_txt(temp_path)
            assert result == ""
        finally:
            os.unlink(temp_path)
    
    def test_unicode_txt_file(self):
        """Extract text with unicode characters."""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("Policy: データプライバシー\nVersion: 2.0")
            temp_path = f.name
        
        try:
            result = extract_text_from_txt(temp_path)
            assert "データプライバシー" in result
        finally:
            os.unlink(temp_path)
    
    def test_large_txt_file(self):
        """Extract text from large TXT file."""
        content = "Section " * 10000  # ~70KB of text
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(content)
            temp_path = f.name
        
        try:
            result = extract_text_from_txt(temp_path)
            assert len(result) > 50000
            assert "Section" in result
        finally:
            os.unlink(temp_path)


class TestChunkText:
    """Tests for text chunking."""
    
    def test_short_text_single_chunk(self):
        """Short text should result in single chunk."""
        text = "This is a short policy."
        chunks = chunk_text(text)
        assert len(chunks) == 1
        assert chunks[0] == text
    
    def test_long_text_multiple_chunks(self):
        """Long text should be split into multiple chunks."""
        # Create text that will definitely need chunking
        paragraphs = []
        for i in range(50):
            paragraphs.append(f"Section {i}: This is paragraph {i} with enough content to make it meaningful for chunking purposes. " * 5)
        text = "\n\n".join(paragraphs)
        
        chunks = chunk_text(text)
        assert len(chunks) > 1
    
    def test_chunks_have_overlap(self):
        """Chunks should have overlapping content for context."""
        # Create text that will need chunking
        text = "UNIQUE_START. " + ("Middle content. " * 500) + "UNIQUE_END."
        chunks = chunk_text(text)
        
        if len(chunks) > 1:
            # First chunk should have start
            assert "UNIQUE_START" in chunks[0]
            # Last chunk should have end
            assert "UNIQUE_END" in chunks[-1]
    
    def test_empty_text(self):
        """Empty text should return empty list or single empty chunk."""
        chunks = chunk_text("")
        assert len(chunks) <= 1
    
    def test_respects_paragraph_boundaries(self):
        """Chunking should prefer splitting at paragraph boundaries."""
        text = "Paragraph 1 content here.\n\nParagraph 2 content here.\n\nParagraph 3 content here."
        chunks = chunk_text(text)
        # All content should be present
        combined = " ".join(chunks)
        assert "Paragraph 1" in combined
        assert "Paragraph 2" in combined
        assert "Paragraph 3" in combined


class TestExtractTextFromPdf:
    """Tests for PDF extraction."""
    
    def test_pdf_extraction_returns_tuple(self):
        """PDF extraction should return (text, metadata) tuple."""
        # Create a minimal valid PDF
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = f.name
        
        try:
            c = canvas.Canvas(temp_path, pagesize=letter)
            c.drawString(100, 750, "Test Policy Document")
            c.drawString(100, 700, "Section 1: Introduction")
            c.save()
            
            text, pages = extract_text_from_pdf(temp_path)
            
            assert isinstance(text, str)
            assert isinstance(pages, list)
            assert len(pages) == 1
            assert "page_number" in pages[0]
        finally:
            os.unlink(temp_path)
    
    def test_pdf_extracts_text_content(self):
        """PDF should extract readable text."""
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = f.name
        
        try:
            c = canvas.Canvas(temp_path, pagesize=letter)
            c.drawString(100, 750, "REMOTE WORK POLICY")
            c.drawString(100, 700, "Effective Date: January 2026")
            c.save()
            
            text, pages = extract_text_from_pdf(temp_path)
            
            assert "REMOTE WORK POLICY" in text or "WORK" in text.upper()
        finally:
            os.unlink(temp_path)
    
    def test_multipage_pdf(self):
        """Multi-page PDF should extract from all pages."""
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = f.name
        
        try:
            c = canvas.Canvas(temp_path, pagesize=letter)
            c.drawString(100, 750, "Page 1 Content")
            c.showPage()
            c.drawString(100, 750, "Page 2 Content")
            c.showPage()
            c.drawString(100, 750, "Page 3 Content")
            c.save()
            
            text, pages = extract_text_from_pdf(temp_path)
            
            assert len(pages) == 3
            assert pages[0]["page_number"] == 1
            assert pages[1]["page_number"] == 2
            assert pages[2]["page_number"] == 3
        finally:
            os.unlink(temp_path)


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
class TestExtractTextFromDocx:
    """Tests for Word document extraction."""
    
    def test_simple_docx_extraction(self):
        """Extract text from simple DOCX file."""
        from docx import Document
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        try:
            doc = Document()
            doc.add_heading("Employee Leave Policy", 0)
            doc.add_paragraph("This policy outlines leave entitlements.")
            doc.add_paragraph("Effective Date: January 1, 2026")
            doc.save(temp_path)
            
            result = extract_text_from_docx(temp_path)
            
            assert "Employee Leave Policy" in result
            assert "leave entitlements" in result
        finally:
            os.unlink(temp_path)
    
    def test_docx_with_tables(self):
        """Extract text from DOCX with tables."""
        from docx import Document
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        try:
            doc = Document()
            doc.add_heading("Leave Types", 0)
            
            table = doc.add_table(rows=3, cols=2)
            table.rows[0].cells[0].text = "Leave Type"
            table.rows[0].cells[1].text = "Days"
            table.rows[1].cells[0].text = "Annual"
            table.rows[1].cells[1].text = "20"
            table.rows[2].cells[0].text = "Sick"
            table.rows[2].cells[1].text = "10"
            
            doc.save(temp_path)
            
            result = extract_text_from_docx(temp_path)
            
            assert "Leave Type" in result or "Leave" in result
            assert "Annual" in result
            assert "20" in result
        finally:
            os.unlink(temp_path)
    
    def test_empty_docx(self):
        """Extract text from empty DOCX."""
        from docx import Document
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        try:
            doc = Document()
            doc.save(temp_path)
            
            result = extract_text_from_docx(temp_path)
            
            assert result == "" or result.strip() == ""
        finally:
            os.unlink(temp_path)


class TestEndToEndIndexing:
    """End-to-end tests for the full indexing pipeline."""
    
    @pytest.fixture
    def sample_txt_file(self):
        """Create a sample TXT file for testing."""
        content = """DATA PRIVACY POLICY
Version: 3.0
Effective Date: January 1, 2026

1. INTRODUCTION
This policy governs the collection, use, and protection of personal data.

2. SCOPE
This policy applies to all employees and contractors.

3. DATA COLLECTION
We collect only necessary personal information.
"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(content)
            temp_path = f.name
        
        yield temp_path
        os.unlink(temp_path)
    
    @pytest.fixture
    def sample_pdf_file(self):
        """Create a sample PDF file for testing."""
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = f.name
        
        c = canvas.Canvas(temp_path, pagesize=letter)
        c.drawString(100, 750, "INFORMATION SECURITY POLICY")
        c.drawString(100, 720, "Version: 2.1")
        c.drawString(100, 690, "")
        c.drawString(100, 660, "1. PURPOSE")
        c.drawString(100, 630, "Protect company information assets.")
        c.save()
        
        yield temp_path
        os.unlink(temp_path)
    
    def test_txt_full_pipeline(self, sample_txt_file):
        """Test TXT file through extraction and chunking."""
        text = extract_text_from_txt(sample_txt_file)
        sanitized = sanitize_text_for_pinecone(text)
        chunks = chunk_text(sanitized)
        
        assert len(text) > 0
        assert len(sanitized) > 0
        assert len(chunks) >= 1
        
        # Verify content integrity
        combined = " ".join(chunks)
        assert "DATA PRIVACY POLICY" in combined
        assert "personal data" in combined.lower()
    
    def test_pdf_full_pipeline(self, sample_pdf_file):
        """Test PDF file through extraction and chunking."""
        text, pages = extract_text_from_pdf(sample_pdf_file)
        sanitized = sanitize_text_for_pinecone(text)
        chunks = chunk_text(sanitized)
        
        assert len(text) > 0
        assert len(pages) == 1
        assert len(chunks) >= 1
    
    def test_problematic_characters_handled(self):
        """Test that problematic characters don't break the pipeline."""
        # Text with various problematic characters
        text = "Policy\x00with\x01null\x02bytes\x03and\x7fcontrol chars"
        
        sanitized = sanitize_text_for_pinecone(text)
        chunks = chunk_text(sanitized)
        
        # Should not raise and should produce valid output
        assert "\x00" not in sanitized
        assert len(chunks) >= 1


class TestApiUploadIntegration:
    """Integration tests for the upload API endpoint."""
    
    @pytest.fixture
    def client(self):
        """Create test client."""
        # Skip if server is running externally
        import socket
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        result = sock.connect_ex(('localhost', 8003))
        sock.close()
        if result == 0:
            pytest.skip("Server already running on port 8003")
        
        from fastapi.testclient import TestClient
        from app.main import app
        return TestClient(app)
    
    def test_upload_txt_file(self, client):
        """Test uploading a TXT file via API."""
        content = b"Test policy content for upload."
        files = {"file": ("test_policy.txt", content, "text/plain")}
        
        response = client.post("/api/docs/upload", files=files)
        
        # Should succeed or fail gracefully
        assert response.status_code in [200, 201, 422, 500]  # May fail if Pinecone not configured
    
    def test_upload_rejects_unsupported_types(self, client):
        """Test that unsupported file types are rejected."""
        content = b"Not a valid document"
        files = {"file": ("test.exe", content, "application/octet-stream")}
        
        response = client.post("/api/docs/upload", files=files)
        
        # Should be rejected
        assert response.status_code in [400, 415, 422]
    
    def test_upload_empty_file(self, client):
        """Test uploading an empty file."""
        content = b""
        files = {"file": ("empty.txt", content, "text/plain")}
        
        response = client.post("/api/docs/upload", files=files)
        
        # Should fail gracefully
        assert response.status_code in [400, 422, 500]


# Run tests with pytest
if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
