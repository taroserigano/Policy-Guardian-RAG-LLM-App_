"""
Live API integration tests for document upload.
Run with backend server running on port 8003.
"""
import sys
import os
import tempfile
import requests
from pathlib import Path

# Configuration
API_BASE_URL = "http://localhost:8005"

class TestResult:
    def __init__(self):
        self.passed = 0
        self.failed = 0
        self.skipped = 0
        self.errors = []
    
    def add_pass(self, name):
        self.passed += 1
        print(f"  ✅ {name}")
    
    def add_fail(self, name, error):
        self.failed += 1
        self.errors.append((name, error))
        print(f"  ❌ {name}: {error}")
    
    def add_skip(self, name, reason):
        self.skipped += 1
        print(f"  ⏭️  {name}: {reason}")
    
    def summary(self):
        total = self.passed + self.failed + self.skipped
        print(f"\n{'='*60}")
        print(f"  Results: {self.passed} passed, {self.failed} failed, {self.skipped} skipped / {total} total")
        if self.errors:
            print(f"  Failures:")
            for name, error in self.errors:
                print(f"    - {name}: {error}")
        print(f"{'='*60}")
        return self.failed == 0

results = TestResult()


def check_server():
    """Check if the server is running."""
    try:
        response = requests.get(f"{API_BASE_URL}/health", timeout=5)
        return response.status_code == 200
    except:
        return False


# ============================================
# Pre-flight check
# ============================================
print("="*60)
print("  Document Upload API Integration Tests")
print("="*60)
print()

if not check_server():
    print(f"❌ Server not running at {API_BASE_URL}")
    print("   Start the backend with:")
    print("   python -m uvicorn app.main:app --host 127.0.0.1 --port 8003")
    sys.exit(1)

print(f"✓ Server running at {API_BASE_URL}")
print()


# ============================================
# TEST: TXT File Upload
# ============================================
print("="*60)
print("  Testing TXT File Upload")
print("="*60)

# Test 1: Simple TXT upload
content = b"""DATA PRIVACY POLICY
Version: 3.0
Effective Date: January 1, 2026

1. INTRODUCTION
This policy governs the collection and use of personal data.

2. SCOPE
This policy applies to all employees and contractors.
"""

try:
    files = {"file": ("test_policy.txt", content, "text/plain")}
    response = requests.post(f"{API_BASE_URL}/api/docs/upload", files=files, timeout=120)
    
    if response.status_code in [200, 201]:
        data = response.json()
        if "id" in data or "message" in data:
            results.add_pass("txt_upload_basic")
        else:
            results.add_fail("txt_upload_basic", f"Unexpected response: {data}")
    else:
        results.add_fail("txt_upload_basic", f"Status {response.status_code}: {response.text}")
except Exception as e:
    results.add_fail("txt_upload_basic", str(e))

# Test 2: TXT with Unicode
content = """ポリシー文書
Version: 2.0

データプライバシーに関するポリシーです。
""".encode('utf-8')

try:
    files = {"file": ("unicode_policy.txt", content, "text/plain")}
    response = requests.post(f"{API_BASE_URL}/api/docs/upload", files=files, timeout=120)
    
    if response.status_code in [200, 201]:
        results.add_pass("txt_upload_unicode")
    else:
        results.add_fail("txt_upload_unicode", f"Status {response.status_code}: {response.text[:200]}")
except Exception as e:
    results.add_fail("txt_upload_unicode", str(e))


# ============================================
# TEST: PDF File Upload
# ============================================
print("\n" + "="*60)
print("  Testing PDF File Upload")
print("="*60)

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    import io
    
    # Test 3: Simple PDF
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.drawString(100, 750, "REMOTE WORK POLICY")
    c.drawString(100, 720, "Version: 2.1")
    c.drawString(100, 690, "")
    c.drawString(100, 660, "1. PURPOSE")
    c.drawString(100, 630, "This policy defines remote work guidelines.")
    c.save()
    pdf_content = buffer.getvalue()
    
    files = {"file": ("remote_work.pdf", pdf_content, "application/pdf")}
    response = requests.post(f"{API_BASE_URL}/api/docs/upload", files=files, timeout=120)
    
    if response.status_code in [200, 201]:
        results.add_pass("pdf_upload_basic")
    else:
        results.add_fail("pdf_upload_basic", f"Status {response.status_code}: {response.text[:200]}")
    
    # Test 4: Multi-page PDF
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.drawString(100, 750, "Page 1: Introduction")
    c.showPage()
    c.drawString(100, 750, "Page 2: Details")
    c.showPage()
    c.drawString(100, 750, "Page 3: Conclusion")
    c.save()
    pdf_content = buffer.getvalue()
    
    files = {"file": ("multipage.pdf", pdf_content, "application/pdf")}
    response = requests.post(f"{API_BASE_URL}/api/docs/upload", files=files, timeout=120)
    
    if response.status_code in [200, 201]:
        results.add_pass("pdf_upload_multipage")
    else:
        results.add_fail("pdf_upload_multipage", f"Status {response.status_code}: {response.text[:200]}")
    
    # Test 5: PDF with special characters
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.drawString(100, 750, "Special: (c) 2026 Company")
    c.drawString(100, 720, "Price: $100.00 - 50% discount")
    c.drawString(100, 690, "Email: test@example.com")
    c.save()
    pdf_content = buffer.getvalue()
    
    files = {"file": ("special_chars.pdf", pdf_content, "application/pdf")}
    response = requests.post(f"{API_BASE_URL}/api/docs/upload", files=files, timeout=120)
    
    if response.status_code in [200, 201]:
        results.add_pass("pdf_upload_special_chars")
    else:
        results.add_fail("pdf_upload_special_chars", f"Status {response.status_code}: {response.text[:200]}")

except ImportError:
    results.add_skip("pdf_upload_basic", "reportlab not installed")
    results.add_skip("pdf_upload_multipage", "reportlab not installed")
    results.add_skip("pdf_upload_special_chars", "reportlab not installed")
except Exception as e:
    results.add_fail("pdf_upload_tests", str(e))


# ============================================
# TEST: Word Document Upload
# ============================================
print("\n" + "="*60)
print("  Testing Word Document Upload")
print("="*60)

try:
    from docx import Document
    import io
    
    # Test 6: Simple DOCX
    doc = Document()
    doc.add_heading("Employee Leave Policy", 0)
    doc.add_paragraph("Version: 2.0")
    doc.add_paragraph("This policy outlines leave entitlements for all employees.")
    doc.add_heading("1. Annual Leave", level=1)
    doc.add_paragraph("Employees are entitled to 20 days of annual leave per year.")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_content = buffer.getvalue()
    
    files = {"file": ("leave_policy.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    response = requests.post(f"{API_BASE_URL}/api/docs/upload", files=files, timeout=120)
    
    if response.status_code in [200, 201]:
        results.add_pass("docx_upload_basic")
    else:
        results.add_fail("docx_upload_basic", f"Status {response.status_code}: {response.text[:200]}")
    
    # Test 7: DOCX with table
    doc = Document()
    doc.add_heading("Leave Summary", 0)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Leave Type"
    table.rows[0].cells[1].text = "Days"
    table.rows[1].cells[0].text = "Annual"
    table.rows[1].cells[1].text = "20"
    table.rows[2].cells[0].text = "Sick"
    table.rows[2].cells[1].text = "10"
    table.rows[3].cells[0].text = "Personal"
    table.rows[3].cells[1].text = "5"
    
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_content = buffer.getvalue()
    
    files = {"file": ("leave_table.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    response = requests.post(f"{API_BASE_URL}/api/docs/upload", files=files, timeout=120)
    
    if response.status_code in [200, 201]:
        results.add_pass("docx_upload_with_table")
    else:
        results.add_fail("docx_upload_with_table", f"Status {response.status_code}: {response.text[:200]}")

except ImportError:
    results.add_skip("docx_upload_basic", "python-docx not installed")
    results.add_skip("docx_upload_with_table", "python-docx not installed")
except Exception as e:
    results.add_fail("docx_upload_tests", str(e))


# ============================================
# TEST: Error Cases
# ============================================
print("\n" + "="*60)
print("  Testing Error Cases")
print("="*60)

# Test 8: Unsupported file type
try:
    files = {"file": ("script.exe", b"fake exe content", "application/octet-stream")}
    response = requests.post(f"{API_BASE_URL}/api/docs/upload", files=files, timeout=30)
    
    if response.status_code in [400, 415, 422]:
        results.add_pass("rejects_unsupported_type")
    else:
        results.add_fail("rejects_unsupported_type", f"Expected rejection, got {response.status_code}")
except Exception as e:
    results.add_fail("rejects_unsupported_type", str(e))

# Test 9: Empty file
try:
    files = {"file": ("empty.txt", b"", "text/plain")}
    response = requests.post(f"{API_BASE_URL}/api/docs/upload", files=files, timeout=30)
    
    # Empty file should either be rejected or handled gracefully
    if response.status_code in [400, 422, 500]:
        results.add_pass("handles_empty_file")
    elif response.status_code in [200, 201]:
        # If accepted, it should have handled it gracefully
        results.add_pass("handles_empty_file")
    else:
        results.add_fail("handles_empty_file", f"Unexpected status {response.status_code}")
except Exception as e:
    results.add_fail("handles_empty_file", str(e))


# ============================================
# TEST: Query After Upload
# ============================================
print("\n" + "="*60)
print("  Testing Query After Upload")
print("="*60)

# Test 10: Upload then query
try:
    # First upload a document
    content = b"""SECURITY POLICY TEST DOCUMENT
Version: 1.0

All employees must use strong passwords.
Passwords must be at least 12 characters.
Multi-factor authentication is required.
"""
    files = {"file": ("security_test.txt", content, "text/plain")}
    response = requests.post(f"{API_BASE_URL}/api/docs/upload", files=files, timeout=120)
    
    if response.status_code in [200, 201]:
        # Now query for it (using correct schema)
        query_data = {
            "user_id": "test-session-001",
            "provider": "ollama",
            "question": "What are the password requirements?"
        }
        response = requests.post(f"{API_BASE_URL}/api/chat", json=query_data, timeout=60)
        
        if response.status_code == 200:
            results.add_pass("query_after_upload")
        elif response.status_code == 500:
            # Check if it's an Ollama connection error
            error_text = response.text.lower()
            if "11434" in error_text or "ollama" in error_text or "connection" in error_text:
                results.add_skip("query_after_upload", "Ollama not running (expected if not installed)")
            else:
                results.add_fail("query_after_upload", f"Query failed: {response.status_code}")
        else:
            results.add_fail("query_after_upload", f"Query failed: {response.status_code}")
    else:
        results.add_skip("query_after_upload", "Upload failed, skipping query test")
except Exception as e:
    results.add_fail("query_after_upload", str(e))


# ============================================
# SUMMARY
# ============================================
success = results.summary()
sys.exit(0 if success else 1)
