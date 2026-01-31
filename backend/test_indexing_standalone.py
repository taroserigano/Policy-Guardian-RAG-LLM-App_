"""
Standalone tests for document indexing - no external dependencies required.
Run with: python test_indexing_standalone.py
"""
import sys
import os
import tempfile
from pathlib import Path

# Minimal test framework
class TestResult:
    def __init__(self):
        self.passed = 0
        self.failed = 0
        self.errors = []
    
    def add_pass(self, name):
        self.passed += 1
        print(f"  ✅ {name}")
    
    def add_fail(self, name, error):
        self.failed += 1
        self.errors.append((name, error))
        print(f"  ❌ {name}: {error}")
    
    def summary(self):
        total = self.passed + self.failed
        print(f"\n{'='*60}")
        print(f"  Results: {self.passed}/{total} passed")
        if self.errors:
            print(f"  Failures:")
            for name, error in self.errors:
                print(f"    - {name}: {error}")
        print(f"{'='*60}")
        return self.failed == 0

results = TestResult()


# ============================================
# TEST: sanitize_text_for_pinecone
# ============================================
print("\n" + "="*60)
print("  Testing sanitize_text_for_pinecone")
print("="*60)

# Inline implementation to test without imports
MAX_METADATA_SIZE = 40000

def sanitize_text_for_pinecone(text):
    """Sanitize text for Pinecone storage."""
    if not text:
        return ""
    
    sanitized = ""
    for char in text:
        if char in ('\n', '\r', '\t') or (ord(char) >= 32 and ord(char) != 127):
            sanitized += char
    
    sanitized = sanitized.encode('utf-8', errors='replace').decode('utf-8')
    
    if len(sanitized.encode('utf-8')) > MAX_METADATA_SIZE:
        while len(sanitized.encode('utf-8')) > MAX_METADATA_SIZE:
            sanitized = sanitized[:-100]
        sanitized += "... [truncated]"
    
    return sanitized.strip()


# Test 1: Empty string
try:
    assert sanitize_text_for_pinecone("") == ""
    results.add_pass("empty_string")
except AssertionError as e:
    results.add_fail("empty_string", str(e))

# Test 2: None input
try:
    assert sanitize_text_for_pinecone(None) == ""
    results.add_pass("none_input")
except AssertionError as e:
    results.add_fail("none_input", str(e))

# Test 3: Normal text
try:
    text = "This is a normal policy document."
    assert sanitize_text_for_pinecone(text) == text
    results.add_pass("normal_text")
except AssertionError as e:
    results.add_fail("normal_text", str(e))

# Test 4: Preserves newlines
try:
    text = "Line 1\nLine 2\nLine 3"
    result = sanitize_text_for_pinecone(text)
    assert "\n" in result
    results.add_pass("preserves_newlines")
except AssertionError as e:
    results.add_fail("preserves_newlines", str(e))

# Test 5: Preserves tabs
try:
    text = "Col1\tCol2\tCol3"
    result = sanitize_text_for_pinecone(text)
    assert "\t" in result
    results.add_pass("preserves_tabs")
except AssertionError as e:
    results.add_fail("preserves_tabs", str(e))

# Test 6: Removes null bytes
try:
    text = "Hello\x00World"
    result = sanitize_text_for_pinecone(text)
    assert "\x00" not in result
    assert "HelloWorld" in result
    results.add_pass("removes_null_bytes")
except AssertionError as e:
    results.add_fail("removes_null_bytes", str(e))

# Test 7: Removes control characters
try:
    text = "Hello\x01\x02\x03World"
    result = sanitize_text_for_pinecone(text)
    assert "\x01" not in result
    assert "\x02" not in result
    assert "HelloWorld" in result
    results.add_pass("removes_control_chars")
except AssertionError as e:
    results.add_fail("removes_control_chars", str(e))

# Test 8: Removes DEL character
try:
    text = "Hello\x7fWorld"
    result = sanitize_text_for_pinecone(text)
    assert "\x7f" not in result
    results.add_pass("removes_del_char")
except AssertionError as e:
    results.add_fail("removes_del_char", str(e))

# Test 9: Handles Unicode
try:
    text = "日本語テスト émojis 🎉 symbols ™®©"
    result = sanitize_text_for_pinecone(text)
    assert "日本語" in result
    assert "🎉" in result
    results.add_pass("handles_unicode")
except AssertionError as e:
    results.add_fail("handles_unicode", str(e))

# Test 10: Truncates long text
try:
    long_text = "A" * 50000
    result = sanitize_text_for_pinecone(long_text)
    assert len(result.encode('utf-8')) <= MAX_METADATA_SIZE + 100
    assert "... [truncated]" in result
    results.add_pass("truncates_long_text")
except AssertionError as e:
    results.add_fail("truncates_long_text", str(e))

# Test 11: Strips whitespace
try:
    text = "   Hello World   "
    result = sanitize_text_for_pinecone(text)
    assert result == "Hello World"
    results.add_pass("strips_whitespace")
except AssertionError as e:
    results.add_fail("strips_whitespace", str(e))


# ============================================
# TEST: PDF Text Extraction
# ============================================
print("\n" + "="*60)
print("  Testing PDF Text Extraction")
print("="*60)

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    import pypdf
    
    # Test 12: Create and extract from PDF
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
        temp_path = f.name
    
    try:
        c = canvas.Canvas(temp_path, pagesize=letter)
        c.drawString(100, 750, "REMOTE WORK POLICY")
        c.drawString(100, 720, "Version: 2.1")
        c.drawString(100, 690, "Effective Date: January 2026")
        c.save()
        
        with open(temp_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text() or ""
                text += page_text
        
        assert len(text) > 0
        results.add_pass("pdf_extraction_basic")
    finally:
        os.unlink(temp_path)
    
    # Test 13: Multi-page PDF
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
        temp_path = f.name
    
    try:
        c = canvas.Canvas(temp_path, pagesize=letter)
        c.drawString(100, 750, "Page 1")
        c.showPage()
        c.drawString(100, 750, "Page 2")
        c.showPage()
        c.drawString(100, 750, "Page 3")
        c.save()
        
        with open(temp_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            assert len(pdf_reader.pages) == 3
        
        results.add_pass("pdf_multipage")
    finally:
        os.unlink(temp_path)
    
    # Test 14: PDF with special characters
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
        temp_path = f.name
    
    try:
        c = canvas.Canvas(temp_path, pagesize=letter)
        c.drawString(100, 750, "Special: Copyright (c) 2026")
        c.drawString(100, 720, "Price: $100.00")
        c.save()
        
        with open(temp_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            text = pdf_reader.pages[0].extract_text() or ""
            sanitized = sanitize_text_for_pinecone(text)
        
        assert len(sanitized) > 0
        results.add_pass("pdf_special_chars")
    finally:
        os.unlink(temp_path)
        
except ImportError as e:
    results.add_fail("pdf_tests", f"Missing library: {e}")


# ============================================
# TEST: Word Document Extraction
# ============================================
print("\n" + "="*60)
print("  Testing Word Document Extraction")
print("="*60)

try:
    from docx import Document
    
    # Test 15: Simple DOCX
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = f.name
    
    try:
        doc = Document()
        doc.add_heading("Employee Leave Policy", 0)
        doc.add_paragraph("This policy outlines leave entitlements.")
        doc.save(temp_path)
        
        doc = Document(temp_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        assert "Employee Leave Policy" in text
        assert "leave entitlements" in text
        results.add_pass("docx_basic_extraction")
    finally:
        os.unlink(temp_path)
    
    # Test 16: DOCX with tables
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = f.name
    
    try:
        doc = Document()
        doc.add_heading("Leave Types", 0)
        
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "Leave Type"
        table.rows[0].cells[1].text = "Days"
        table.rows[1].cells[0].text = "Annual"
        table.rows[1].cells[1].text = "20"
        table.rows[2].cells[0].text = "Sick"
        table.rows[2].cells[1].text = "10"
        doc.save(temp_path)
        
        # Extract including tables
        doc = Document(temp_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        text = "\n".join(paragraphs)
        assert "Leave Type" in text or "Annual" in text
        results.add_pass("docx_table_extraction")
    finally:
        os.unlink(temp_path)
    
    # Test 17: Empty DOCX
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = f.name
    
    try:
        doc = Document()
        doc.save(temp_path)
        
        doc = Document(temp_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        assert text.strip() == ""
        results.add_pass("docx_empty")
    finally:
        os.unlink(temp_path)

except ImportError as e:
    results.add_fail("docx_tests", f"Missing library: {e}")


# ============================================
# TEST: TXT File Extraction
# ============================================
print("\n" + "="*60)
print("  Testing TXT File Extraction")
print("="*60)

# Test 18: Simple TXT
with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
    f.write("This is a test policy.\nSecond line.")
    temp_path = f.name

try:
    with open(temp_path, 'r', encoding='utf-8') as f:
        text = f.read()
    assert "test policy" in text
    assert "Second line" in text
    results.add_pass("txt_basic")
finally:
    os.unlink(temp_path)

# Test 19: TXT with Unicode
with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
    f.write("Policy: データプライバシー\nVersion: 2.0")
    temp_path = f.name

try:
    with open(temp_path, 'r', encoding='utf-8') as f:
        text = f.read()
    assert "データプライバシー" in text
    results.add_pass("txt_unicode")
finally:
    os.unlink(temp_path)

# Test 20: Large TXT
content = "Section content. " * 5000
with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
    f.write(content)
    temp_path = f.name

try:
    with open(temp_path, 'r', encoding='utf-8') as f:
        text = f.read()
    assert len(text) > 50000
    results.add_pass("txt_large")
finally:
    os.unlink(temp_path)


# ============================================
# TEST: Text Chunking
# ============================================
print("\n" + "="*60)
print("  Testing Text Chunking")
print("="*60)

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    # Test 21: Short text single chunk
    text = "This is a short policy."
    chunks = splitter.split_text(text)
    assert len(chunks) == 1
    results.add_pass("chunking_short")
    
    # Test 22: Long text multiple chunks
    paragraphs = []
    for i in range(50):
        paragraphs.append(f"Section {i}: This is paragraph {i} with content. " * 5)
    text = "\n\n".join(paragraphs)
    
    chunks = splitter.split_text(text)
    assert len(chunks) > 1
    results.add_pass("chunking_long")
    
    # Test 23: Empty text
    chunks = splitter.split_text("")
    assert len(chunks) <= 1
    results.add_pass("chunking_empty")

except ImportError as e:
    results.add_fail("chunking_tests", f"Missing library: {e}")


# ============================================
# TEST: Full Pipeline (sanitize + extract + chunk)
# ============================================
print("\n" + "="*60)
print("  Testing Full Pipeline")
print("="*60)

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    # Test 24: TXT pipeline
    content = """DATA PRIVACY POLICY
Version: 3.0

1. INTRODUCTION
This policy governs personal data.

2. SCOPE
Applies to all employees.
"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
        f.write(content)
        temp_path = f.name
    
    try:
        with open(temp_path, 'r', encoding='utf-8') as f:
            text = f.read()
        sanitized = sanitize_text_for_pinecone(text)
        chunks = splitter.split_text(sanitized)
        
        assert len(chunks) >= 1
        combined = " ".join(chunks)
        assert "DATA PRIVACY" in combined
        results.add_pass("pipeline_txt")
    finally:
        os.unlink(temp_path)
    
    # Test 25: Pipeline with problematic characters
    text = "Policy\x00with\x01null\x02bytes\x03and\x7fcontrol chars"
    sanitized = sanitize_text_for_pinecone(text)
    chunks = splitter.split_text(sanitized)
    
    assert "\x00" not in sanitized
    assert len(chunks) >= 1
    results.add_pass("pipeline_problematic_chars")

except Exception as e:
    results.add_fail("pipeline_tests", str(e))


# ============================================
# SUMMARY
# ============================================
success = results.summary()
sys.exit(0 if success else 1)
